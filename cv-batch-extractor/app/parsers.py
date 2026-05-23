import logging

import fitz
import pytesseract
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".tiff", ".bmp"}


def pdf_to_text(path: str) -> str:
    doc = fitz.open(path)
    parts: list[str] = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)
        else:
            # Scanned page — render to image and OCR
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img).strip()
            if ocr_text:
                logger.debug("OCR fallback used for page %d of %s", page.number + 1, path)
                parts.append(ocr_text)
    return "\n\n".join(parts)


def docx_to_text(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append("\t".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n".join(parts)


def image_to_text(path: str) -> str:
    img = Image.open(path)
    return pytesseract.image_to_string(img)


def extract_text(path: str) -> str:
    lower = path.lower()
    if lower.endswith(".pdf"):
        return pdf_to_text(path)
    if lower.endswith(".docx") or lower.endswith(".doc"):
        return docx_to_text(path)
    return image_to_text(path)
